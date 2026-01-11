import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_anthropic import ChatAnthropic
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage
from langchain_community.tools.tavily_search import TavilySearchResults
import base64
import pypdf
import docx
import fitz # pymupdf
import pdfplumber

# Disable ChromaDB telemetry to fix PyInstaller issues
os.environ["ANONYMIZED_TELEMETRY"] = "False"

# Initialize settings
# Use current working directory for persistence (works for both dev and PyInstaller EXE)
PERSIST_DIRECTORY = os.path.join(os.getcwd(), "chroma_data")

class RAGEngine:
    def __init__(self, api_key: str, model_name: str = "gemini-1.5-pro", base_url: str = None):
        if not api_key:
            raise ValueError("API Key is required")
        
        # Auto-detect model for OpenAI keys if default is still Gemini
        if api_key.startswith("sk-") and model_name == "gemini-1.5-pro":
             print("DEBUG: Detected OpenAI Key. Switching default model to 'gpt-4o'.")
             model_name = "gpt-4o"

        print(f"DEBUG: RAGEngine initializing with key: {api_key[:5]}...")
        print(f"DEBUG: Persistence Directory: {PERSIST_DIRECTORY}")
        
        self.api_key = api_key
        self.current_model_name = model_name
        self.base_url = base_url
        
        # Embeddings Selection based on Key Format
        try:
            if api_key.startswith("sk-") and not base_url:
                # Likely OpenAI
                print("DEBUG: Using OpenAIEmbeddings")
                self.embeddings = OpenAIEmbeddings(api_key=api_key)
            else:
                # Default to Google (or try Google if generic)
                # Note: This might fail if it's a non-sk key but not Google.
                # Ideally, we should require specific env vars for embeddings.
                print("DEBUG: Using GoogleGenerativeAIEmbeddings")
                self.embeddings = GoogleGenerativeAIEmbeddings(
                    model="models/embedding-001", 
                    google_api_key=api_key
                )
        except Exception as e:
            print(f"Warning: Failed to initialize embeddings: {e}. RAG features may not work.")
            self.embeddings = None

        if self.embeddings:
            try:
                print("DEBUG: Initializing Chroma Vector Store...")
                self.vector_store = Chroma(
                    persist_directory=PERSIST_DIRECTORY,
                    embedding_function=self.embeddings,
                    collection_name="my_knowledge_base"
                )
                print("DEBUG: Chroma initialized successfully.")
            except Exception as e:
                print(f"Error initializing Vector Store (Chroma): {e}. RAG disabled.")
                self.vector_store = None
        else:
            print("DEBUG: No embeddings, RAG disabled.")
            self.vector_store = None
        
        self.llm = self._create_llm(model_name, api_key, base_url)

    def _create_llm(self, model_name: str, api_key: str, base_url: str = None):
        """Creates the LLM instance based on model name and provider config."""
        
        # 1. Google Gemini
        if model_name.lower().startswith("gemini"):
            return ChatGoogleGenerativeAI(
                model=model_name,
                temperature=0,
                google_api_key=api_key
            )
        
        # 2. Native Anthropic (only if no custom base_url is set)
        elif model_name.lower().startswith("claude") and not base_url:
            return ChatAnthropic(
                model=model_name,
                temperature=0,
                api_key=api_key
            )
            
        # 3. OpenAI Compatible (OpenAI, OpenRouter, Groq, DeepSeek, Local)
        else:
            kwargs = {
                "model": model_name,
                "api_key": api_key,
                "request_timeout": 60 # Prevent infinite hanging
            }
            
            # Special handling for reasoning models or models that reject temp=0
            kwargs["temperature"] = 1
            
            # Auto-detect OpenRouter Key
            if api_key.startswith("sk-or-v1") and not base_url:
                print("DEBUG: Detected OpenRouter Key. Setting base_url to 'https://openrouter.ai/api/v1'")
                base_url = "https://openrouter.ai/api/v1"

            if base_url:
                kwargs["base_url"] = base_url
                
            return ChatOpenAI(**kwargs)

    def perform_ocr(self, image_bytes):
        """Uses the current LLM to perform OCR on an image."""
        print("DEBUG: Performing OCR with LLM...")
        try:
            # Encode image to base64
            b64_image = base64.b64encode(image_bytes).decode('utf-8')
            image_data_url = f"data:image/png;base64,{b64_image}"
            
            prompt = "Transcribe the text in this image exactly. Do not add any commentary. Output only the text content."
            
            message = HumanMessage(
                content=[
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": image_data_url}}
                ]
            )
            
            response = self.llm.invoke([message])
            return response.content
        except Exception as e:
            print(f"OCR Failed: {e}")
            return ""

    def ingest_text(self, text: str, source: str = "manual_input", session_id: str = None):
        """Splits and indexes text into the vector store."""
        print(f"DEBUG: ingest_text called for source: {source}. Text length: {len(text)}")
        if self.vector_store is None:
            print("DEBUG: Vector store is None. Skipping ingestion.")
            return 0
            
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)
        print(f"DEBUG: Text split into {len(chunks)} chunks.")
        
        metadata = {"source": source}
        if session_id:
            metadata["session_id"] = session_id

        documents = [Document(page_content=chunk, metadata=metadata) for chunk in chunks]
        
        if documents:
            try:
                print(f"DEBUG: Adding {len(documents)} documents to vector store...")
                self.vector_store.add_documents(documents)
                # Try to force persist if method exists (older versions)
                if hasattr(self.vector_store, 'persist'):
                    print("DEBUG: Persisting vector store...")
                    self.vector_store.persist()
                print("DEBUG: Documents added successfully.")
                return len(documents)
            except Exception as e:
                print(f"ERROR: Failed to add documents to vector store: {e}")
                return 0
        else:
             print("DEBUG: No documents created from text.")
        return 0

    def ingest_file(self, file_path: str, source: str, session_id: str = None):
        """Extracts text from file and indexes it."""
        print(f"DEBUG: ingest_file called for {source} at {file_path}")
        text = ""
        ext = os.path.splitext(source)[1].lower()
        
        try:
            if ext == ".pdf":
                print("DEBUG: Processing PDF with pypdf...")
                try:
                    reader = pypdf.PdfReader(file_path)
                    for i, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                except Exception as e:
                    print(f"DEBUG: pypdf failed: {e}")

                # Fallback 1: PyMuPDF (fitz)
                if not text.strip():
                    print("DEBUG: pypdf yielded empty text. Trying PyMuPDF (fitz)...")
                    doc = None
                    try:
                        doc = fitz.open(file_path)
                        for page in doc:
                            text += page.get_text() + "\n"
                    except Exception as e:
                        print(f"DEBUG: PyMuPDF failed: {e}")
                    finally:
                        if doc:
                            doc.close()
                
                # Fallback 2: pdfplumber
                if not text.strip():
                     print("DEBUG: PyMuPDF yielded empty text. Trying pdfplumber...")
                     try:
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n"
                     except Exception as e:
                        print(f"DEBUG: pdfplumber failed: {e}")

                # Fallback 3: LLM Vision OCR (if text is still empty)
                if not text.strip() and self.llm:
                     print("DEBUG: Text extraction failed. Attempting OCR with LLM Vision...")
                     doc = None
                     try:
                        doc = fitz.open(file_path)
                        for i, page in enumerate(doc):
                            # Limit to first 5 pages to prevent timeouts/high costs
                            if i >= 5:
                                print("DEBUG: Reached page limit (5) for OCR.")
                                break
                                
                            # Render page to image
                            pix = page.get_pixmap()
                            img_bytes = pix.tobytes("png")
                            
                            print(f"DEBUG: OCR Processing page {i+1}/{len(doc)}...")
                            page_text = self.perform_ocr(img_bytes)
                            text += page_text + "\n"
                     except Exception as e:
                        print(f"DEBUG: OCR failed: {e}")
                     finally:
                        if doc:
                            doc.close()

            elif ext == ".docx":
                print("DEBUG: Processing DOCX...")
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == ".txt":
                print("DEBUG: Processing TXT...")
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                print(f"DEBUG: Unsupported file extension: {ext}")
                return 0 # Unsupported file type
            
            print(f"DEBUG: Extracted total text length: {len(text)}")
            if len(text.strip()) == 0:
                 print("DEBUG: Warning - Extracted text is empty! File might be an image/scan.")
                
            return self.ingest_text(text, source, session_id=session_id)
        except Exception as e:
            print(f"Error processing file {source}: {e}")
            raise e

    def list_documents(self):
        """Lists all unique uploaded documents."""
        if self.vector_store is None:
            print("DEBUG: list_documents - vector_store is None")
            return []
        
        try:
            # Get all metadata to find unique sources
            print("DEBUG: list_documents - Fetching from vector_store...")
            data = self.vector_store.get()
            
            if not data or not data['metadatas']:
                print("DEBUG: list_documents - No metadata found in DB.")
                return []
            
            print(f"DEBUG: list_documents - Found {len(data['metadatas'])} chunks.")
            
            sources = set()
            for metadata in data['metadatas']:
                if metadata and 'source' in metadata:
                    sources.add(metadata['source'])
            
            print(f"DEBUG: list_documents - Unique sources: {list(sources)}")
            return list(sources)
        except Exception as e:
            print(f"Error listing documents: {e}")
            return []

    def delete_document(self, source: str):
        """Deletes all chunks associated with a specific source."""
        if self.vector_store is None:
            return False
            
        try:
            # ChromaDB delete by metadata
            # Note: LangChain's Chroma wrapper might not expose delete with where clause easily in all versions,
            # but usually it forwards to the underlying client or we can use _collection.
            
            # Using the underlying collection to be safe and efficient
            self.vector_store._collection.delete(where={"source": source})
            return True
        except Exception as e:
            print(f"Error deleting document {source}: {e}")
            return False

    def get_response(self, query: str, image: str = None, model_name: str = None, base_url: str = None, api_key: str = None, history: list = None, deep_think: bool = False, enable_search: bool = False, search_api_key: str = None, system_instruction: str = None, session_id: str = None):
        """Retrieves context and generates a response."""
        
        # Use provided key or fallback to stored key
        effective_key = api_key or self.api_key

        # Update LLM if model, base_url, or KEY changed
        config_changed = (
            (model_name and model_name != self.current_model_name) or 
            (base_url and base_url != self.base_url) or
            (effective_key != self.api_key) # Check if key changed
        )
        
        if config_changed:
            target_model = model_name or self.current_model_name
            target_base_url = base_url or self.base_url
            
            self.llm = self._create_llm(target_model, effective_key, target_base_url)
            self.current_model_name = target_model
            self.base_url = target_base_url
            self.api_key = effective_key # Update stored key

        # Search Logic
        search_context = ""
        if enable_search and search_api_key:
            try:
                # Set environment variable temporarily for the tool
                os.environ["TAVILY_API_KEY"] = search_api_key
                tool = TavilySearchResults(max_results=3)
                search_results = tool.invoke({"query": query})
                
                # Format results into a context string
                formatted_results = "\n\n--- INTERNET SEARCH RESULTS ---\n"
                for res in search_results:
                    formatted_results += f"Source: {res['url']}\nContent: {res['content']}\n\n"
                formatted_results += "--- END SEARCH RESULTS ---\n\n"
                
                search_context = formatted_results
                print(f"Performed Internet Search. Found {len(search_results)} results.")
            except Exception as e:
                print(f"Internet Search failed: {e}")
                search_context = "\n[Internet Search Attempted but Failed]\n"

        # Add reasoning instruction if Deep Think is enabled
        reasoning_instruction = ""
        if deep_think:
            reasoning_instruction = " Please use a step-by-step reasoning approach and think deeply before providing the final answer."

        # Base System Prompt
        base_system_prompt = system_instruction if system_instruction else "You are a helpful assistant."
        
        # Convert history dicts to LangChain Message objects
        chat_history = []
        if history:
            for msg in history:
                # Skip the very last message if it matches the current query (to avoid duplication)
                if msg.get('role') == 'user':
                    if msg.get('content') == query and msg is history[-1]:
                         continue
                    chat_history.append(HumanMessage(content=msg.get('content', '')))
                elif msg.get('role') == 'model':
                    chat_history.append(AIMessage(content=msg.get('content', '')))

        # If we have a vector store, use RAG. Otherwise just chat.
        if self.vector_store is not None and not image: # Disable RAG if image is present (simplified logic)
            print("DEBUG: Attempting RAG retrieval...")
            try:
                # Apply Session Filter if provided
                search_kwargs = {"k": 5}
                if session_id:
                    print(f"DEBUG: Filtering RAG by session_id: {session_id}")
                    search_kwargs["filter"] = {"session_id": session_id}

                retriever = self.vector_store.as_retriever(search_kwargs=search_kwargs)
                
                # Retrieve documents manually first to debug
                retrieved_docs = retriever.invoke(query)
                print(f"DEBUG: Retrieved {len(retrieved_docs)} documents.")
                for i, doc in enumerate(retrieved_docs):
                    print(f"DEBUG: Doc {i} source: {doc.metadata.get('source', 'unknown')}")
                    # print(f"DEBUG: Doc {i} content preview: {doc.page_content[:50]}...")

                if not retrieved_docs:
                     print("DEBUG: No relevant documents found via RAG.")
                
                system_prompt = (
                    f"{base_system_prompt} "
                    "Use the following pieces of retrieved context to answer "
                    "the question. The context contains information from uploaded files. "
                    "If the context is empty or irrelevant, say that you don't have enough information from the uploaded files. "
                    "Use three sentences maximum and keep the answer concise."
                    + reasoning_instruction +
                    "\n\n"
                    "{context}"
                    + search_context
                )
                
                prompt = ChatPromptTemplate.from_messages(
                    [
                        ("system", system_prompt),
                        MessagesPlaceholder(variable_name="chat_history"),
                        ("human", "{input}"),
                    ]
                )
                
                question_answer_chain = create_stuff_documents_chain(self.llm, prompt)
                rag_chain = create_retrieval_chain(retriever, question_answer_chain)
                
                response = rag_chain.invoke({
                    "input": query,
                    "chat_history": chat_history
                })
                
                # Extract sources
                sources = list(set([doc.metadata.get('source', 'Unknown') for doc in retrieved_docs]))
                
                return {"answer": response["answer"], "sources": sources}
            except Exception as e:
                print(f"RAG Retrieval failed: {e}. Fallback to direct chat.")
                pass
        else:
             print(f"DEBUG: Skipping RAG. VectorStore: {bool(self.vector_store)}, Image: {bool(image)}")
        
        # Fallback to direct chat (or Image Chat)
        print(f"Invoking LLM (DeepThink: {deep_think}, Search: {enable_search}, Image: {bool(image)}) with query: {query[:50]}...")
        try:
            # Construct message content
            if image:
                # Assuming image is base64 data string (e.g. "data:image/png;base64,.....")
                # LangChain expects "image_url" key even for base64 data
                content = [
                    {"type": "text", "text": query},
                    {"type": "image_url", "image_url": {"url": image}} 
                ]
                # Note: This works for OpenAI and some others. 
                # For Gemini via LangChain, it might need specific handling if it doesn't support standard content blocks.
                # However, langchain-google-genai usually handles it.
            else:
                content = query

            # Construct prompt with history and search context
            # System prompt
            system_msg = base_system_prompt + reasoning_instruction + "\n" + search_context
            
            messages = [
                HumanMessage(content=system_msg), # Passing system instruction as HumanMessage first for better compat if system not supported
                *chat_history,
                HumanMessage(content=content)
            ]
            
            response = self.llm.invoke(messages)
            
            print("LLM invocation successful.")
            return {"answer": response.content, "sources": []}
        except Exception as e:
            print(f"LLM invocation failed: {e}")
            raise e
